"""
Enhanced File Processing Service

Comprehensive file processing including text extraction, metadata analysis,
and embedding generation for multiple file types.
"""

import os
import hashlib
import mimetypes
import asyncio
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple, BinaryIO
from pathlib import Path
import tempfile
import magic
import structlog

# File processing imports
import PyPDF2
import pdfplumber
from docx import Document
import openpyxl
from bs4 import BeautifulSoup
from PIL import Image
import json
import csv

# AI and embedding imports
from sentence_transformers import SentenceTransformer
import openai

# Make magic import optional for development
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    # Create a mock magic module for development
    class MockMagic:
        @staticmethod
        def from_file(file_path):
            """Mock magic that returns basic MIME type based on extension"""
            ext = file_path.split('.')[-1].lower()
            mime_types = {
                'pdf': 'application/pdf',
                'txt': 'text/plain',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'jpg': 'image/jpeg',
                'png': 'image/png'
            }
            return mime_types.get(ext, 'application/octet-stream')
    magic = MockMagic()

# Use absolute imports that work when running directly
try:
    # When running as a module
    from database.models import FileRecord, FileStatus, PersonaType, VectorChunk
    from database.vector_store import vector_store
except ImportError:
    # Fallback for relative imports
    from ..database.models import FileRecord, FileStatus, PersonaType, VectorChunk
    from ..database.vector_store import vector_store

logger = structlog.get_logger(__name__)

class FileProcessor:
    """Enhanced file processor with comprehensive format support"""
    
    def __init__(self):
        self.upload_dir = Path(os.getenv("UPLOAD_DIR", "./uploads"))
        self.upload_dir.mkdir(exist_ok=True)
        self.chunk_size = 8192
        self.max_file_size = 2 * 1024 * 1024 * 1024  # 2GB
        
        # Initialize embedding model
        self.embedding_model = None
        self.embedding_model_name = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
        
        # File type processors
        self.processors = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_text,
            'md': self._process_markdown,
            'html': self._process_html,
            'xml': self._process_xml,
            'json': self._process_json,
            'csv': self._process_csv,
            'xlsx': self._process_excel,
            'zip': self._process_archive,
            'tar': self._process_archive,
            '7z': self._process_archive,
            'jpg': self._process_image,
            'jpeg': self._process_image,
            'png': self._process_image,
            'gif': self._process_image,
            'bmp': self._process_image,
            'mp3': self._process_audio,
            'wav': self._process_audio,
            'mp4': self._process_video,
            'avi': self._process_video,
            'py': self._process_code,
            'js': self._process_code,
            'ts': self._process_code,
            'css': self._process_code,
            'sql': self._process_code,
        }
    
    async def initialize(self):
        """Initialize the file processor"""
        try:
            # Load embedding model
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            logger.info(f"Initialized embedding model: {self.embedding_model_name}")
            
            # Initialize vector store
            await vector_store.initialize()
            
            # Create default collection if it doesn't exist
            await vector_store.create_collection(
                "documents",
                self.embedding_model.get_sentence_embedding_dimension(),
                {
                    "file_id": "string",
                    "persona_type": "string",
                    "file_type": "string",
                    "chunk_index": "number"
                }
            )
            
        except Exception as e:
            logger.error("Failed to initialize file processor", error=str(e))
            raise
    
    async def process_upload(self, file_data: BinaryIO, filename: str, 
                           file_size: int, persona_type: PersonaType,
                           metadata: Dict[str, Any]) -> FileRecord:
        """Process uploaded file with comprehensive analysis"""
        
        file_record = FileRecord(
            filename=self._generate_unique_filename(filename),
            original_filename=filename,
            file_size=file_size,
            status=FileStatus.UPLOADING,
            metadata=metadata,
            upload_started=datetime.utcnow()
        )
        
        try:
            # Save file to storage
            storage_path = await self._save_file(file_data, file_record.filename)
            file_record.storage_path = str(storage_path)
            
            # Generate checksum
            file_record.checksum = await self._calculate_checksum(storage_path)
            
            # Detect file type and MIME type
            file_record.mime_type = mimetypes.guess_type(filename)[0]
            file_record.file_type = self._detect_file_type(storage_path)
            
            file_record.status = FileStatus.UPLOADED
            file_record.upload_completed = datetime.utcnow()
            
            # Start background processing
            asyncio.create_task(self._process_file_content(file_record, persona_type))
            
            return file_record
            
        except Exception as e:
            file_record.status = FileStatus.ERROR
            file_record.error_message = str(e)
            logger.error("File upload failed", filename=filename, error=str(e))
            raise
    
    async def _process_file_content(self, file_record: FileRecord, persona_type: PersonaType):
        """Background file content processing"""
        try:
            file_record.status = FileStatus.PROCESSING
            file_record.processing_started = datetime.utcnow()
            
            # Extract content based on file type
            extracted_data = await self._extract_content(file_record)
            
            file_record.extracted_text = extracted_data.get('text', '')
            file_record.extracted_metadata = extracted_data.get('metadata', {})
            
            # Generate embeddings if text content available
            if file_record.extracted_text:
                await self._generate_embeddings(file_record, persona_type)
            
            file_record.status = FileStatus.COMPLETED
            file_record.processing_completed = datetime.utcnow()
            file_record.processing_progress = 1.0
            
            logger.info("File processing completed", 
                       filename=file_record.original_filename,
                       chunks=file_record.chunk_count)
            
        except Exception as e:
            file_record.status = FileStatus.ERROR
            file_record.error_message = str(e)
            logger.error("File processing failed", 
                        filename=file_record.original_filename, 
                        error=str(e))
    
    async def _extract_content(self, file_record: FileRecord) -> Dict[str, Any]:
        """Extract content from file based on type"""
        file_path = Path(file_record.storage_path)
        file_type = file_record.file_type.lower()
        
        if file_type in self.processors:
            return await self.processors[file_type](file_path)
        else:
            # Default text extraction for unknown types
            return await self._process_text(file_path)
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files"""
        text = ""
        metadata = {}
        
        try:
            # Try pdfplumber first for better text extraction
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                text = "\n\n".join(text_parts)
                metadata.update({
                    'pages': len(pdf.pages),
                    'producer': pdf.metadata.get('Producer', ''),
                    'creator': pdf.metadata.get('Creator', ''),
                    'title': pdf.metadata.get('Title', ''),
                    'author': pdf.metadata.get('Author', ''),
                })
        
        except Exception as e:
            logger.warning("pdfplumber failed, trying PyPDF2", error=str(e))
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page in pdf_reader.pages:
                        text_parts.append(page.extract_text())
                    
                    text = "\n\n".join(text_parts)
                    metadata.update({
                        'pages': len(pdf_reader.pages),
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                    })
            
            except Exception as e2:
                logger.error("PDF processing failed", error=str(e2))
                raise
        
        return {
            'text': text,
            'metadata': metadata,
            'content_type': 'document'
        }
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process DOCX files"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'images': len(doc.inline_shapes)
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'document'
            }
            
        except Exception as e:
            logger.error("DOCX processing failed", error=str(e))
            raise
    
    async def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                raise Exception("Could not decode text file")
            
            # Basic text analysis
            lines = text.split('\n')
            words = text.split()
            
            metadata = {
                'lines': len(lines),
                'words': len(words),
                'characters': len(text),
                'encoding': encoding
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'text'
            }
            
        except Exception as e:
            logger.error("Text processing failed", error=str(e))
            raise
    
    async def _process_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Process Markdown files"""
        result = await self._process_text(file_path)
        result['content_type'] = 'markdown'
        
        # Count markdown elements
        text = result['text']
        result['metadata'].update({
            'headers': text.count('#'),
            'links': text.count(']('),
            'code_blocks': text.count('```'),
            'bold_text': text.count('**'),
            'italic_text': text.count('*')
        })
        
        return result
    
    async def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Process HTML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract text content
            text = soup.get_text(separator='\n', strip=True)
            
            # Extract metadata
            title = soup.find('title')
            meta_tags = soup.find_all('meta')
            
            metadata = {
                'title': title.get_text() if title else '',
                'links': len(soup.find_all('a')),
                'images': len(soup.find_all('img')),
                'headings': len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                'meta_tags': len(meta_tags)
            }
            
            # Extract meta tag information
            for meta in meta_tags:
                name = meta.get('name') or meta.get('property')
                content = meta.get('content')
                if name and content:
                    metadata[f'meta_{name}'] = content
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'html'
            }
            
        except Exception as e:
            logger.error("HTML processing failed", error=str(e))
            raise
    
    async def _process_xml(self, file_path: Path) -> Dict[str, Any]:
        """Process XML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'xml')
            
            # Extract text content
            text = soup.get_text(separator='\n', strip=True)
            
            # Count XML elements
            all_tags = soup.find_all()
            tag_counts = {}
            for tag in all_tags:
                tag_name = tag.name
                if tag_name:
                    tag_counts[tag_name] = tag_counts.get(tag_name, 0) + 1
            
            # Basic XML analysis
            metadata = {
                'total_elements': len(all_tags),
                'unique_tags': len(tag_counts),
                'tag_counts': tag_counts,
                'root_element': soup.find().name if soup.find() else 'unknown',
                'namespaces': len(set(tag.prefix for tag in all_tags if tag.prefix))
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'xml'
            }
            
        except Exception as e:
            logger.error("XML processing failed", error=str(e))
            # Fallback to text processing
            return await self._process_text(file_path)
    
    async def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert to readable text
            text = json.dumps(data, indent=2)
            
            # Analyze structure
            def count_elements(obj, counts=None):
                if counts is None:
                    counts = {'objects': 0, 'arrays': 0, 'primitives': 0}
                
                if isinstance(obj, dict):
                    counts['objects'] += 1
                    for value in obj.values():
                        count_elements(value, counts)
                elif isinstance(obj, list):
                    counts['arrays'] += 1
                    for item in obj:
                        count_elements(item, counts)
                else:
                    counts['primitives'] += 1
                
                return counts
            
            structure = count_elements(data)
            
            metadata = {
                'size_bytes': len(text),
                'objects': structure['objects'],
                'arrays': structure['arrays'],
                'primitives': structure['primitives'],
                'top_level_keys': list(data.keys()) if isinstance(data, dict) else []
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'structured_data'
            }
            
        except Exception as e:
            logger.error("JSON processing failed", error=str(e))
            raise
    
    async def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            rows = []
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                headers = next(reader, [])
                for row in reader:
                    rows.append(row)
            
            # Convert to text representation
            text = f"Headers: {', '.join(headers)}\n\n"
            for i, row in enumerate(rows[:10]):  # First 10 rows
                text += f"Row {i+1}: {', '.join(row)}\n"
            
            if len(rows) > 10:
                text += f"\n... and {len(rows) - 10} more rows"
            
            metadata = {
                'rows': len(rows),
                'columns': len(headers),
                'headers': headers,
                'sample_data': rows[:5]  # First 5 rows as sample
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'tabular_data'
            }
            
        except Exception as e:
            logger.error("CSV processing failed", error=str(e))
            raise
    
    async def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            
            text_parts = []
            metadata = {
                'sheets': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            }
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                
                # Get data from sheet
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([str(cell) if cell is not None else '' for cell in row])
                
                if data:
                    # Add headers if they exist
                    if data[0]:
                        text_parts.append(f"Headers: {', '.join(data[0])}")
                    
                    # Add sample data
                    for i, row in enumerate(data[1:6]):  # First 5 data rows
                        text_parts.append(f"Row {i+1}: {', '.join(row)}")
                    
                    if len(data) > 6:
                        text_parts.append(f"... and {len(data) - 6} more rows")
                
                metadata[f'sheet_{sheet_name}'] = {
                    'rows': len(data),
                    'columns': len(data[0]) if data else 0
                }
            
            text = "\n".join(text_parts)
            
            return {
                'text': text,
                'metadata': metadata,
                'content_type': 'spreadsheet'
            }
            
        except Exception as e:
            logger.error("Excel processing failed", error=str(e))
            raise
    
    async def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image files"""
        try:
            with Image.open(file_path) as img:
                # Basic image analysis
                metadata = {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                    'size_pixels': img.width * img.height
                }
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        metadata['exif'] = {str(k): str(v) for k, v in exif.items()}
                
                # Generate description text
                text = f"Image file: {file_path.name}\n"
                text += f"Dimensions: {img.width}x{img.height} pixels\n"
                text += f"Format: {img.format}\n"
                text += f"Color mode: {img.mode}"
                
                return {
                    'text': text,
                    'metadata': metadata,
                    'content_type': 'image'
                }
            
        except Exception as e:
            logger.error("Image processing failed", error=str(e))
            raise
    
    async def _process_code(self, file_path: Path) -> Dict[str, Any]:
        """Process code files"""
        result = await self._process_text(file_path)
        result['content_type'] = 'code'
        
        # Analyze code structure
        text = result['text']
        extension = file_path.suffix.lower()
        
        # Basic code analysis
        code_metrics = {
            'functions': 0,
            'classes': 0,
            'imports': 0,
            'comments': 0
        }
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if extension in ['.py']:
                if line.startswith('def '):
                    code_metrics['functions'] += 1
                elif line.startswith('class '):
                    code_metrics['classes'] += 1
                elif line.startswith('import ') or line.startswith('from '):
                    code_metrics['imports'] += 1
                elif line.startswith('#'):
                    code_metrics['comments'] += 1
            elif extension in ['.js', '.ts']:
                if 'function ' in line or '=>' in line:
                    code_metrics['functions'] += 1
                elif line.startswith('class '):
                    code_metrics['classes'] += 1
                elif line.startswith('import ') or line.startswith('require('):
                    code_metrics['imports'] += 1
                elif line.startswith('//') or line.startswith('/*'):
                    code_metrics['comments'] += 1
        
        result['metadata'].update(code_metrics)
        result['metadata']['language'] = extension
        
        return result
    
    async def _process_archive(self, file_path: Path) -> Dict[str, Any]:
        """Process archive files (ZIP, TAR, etc.)"""
        # For now, just return basic metadata
        # Full archive extraction would require security considerations
        
        file_size = file_path.stat().st_size
        
        metadata = {
            'archive_type': file_path.suffix.lower(),
            'size_bytes': file_size,
            'processing_note': 'Archive content extraction not implemented for security'
        }
        
        text = f"Archive file: {file_path.name}\n"
        text += f"Type: {file_path.suffix.upper()}\n"
        text += f"Size: {file_size} bytes"
        
        return {
            'text': text,
            'metadata': metadata,
            'content_type': 'archive'
        }
    
    async def _process_audio(self, file_path: Path) -> Dict[str, Any]:
        """Process audio files"""
        # Basic audio file analysis
        file_size = file_path.stat().st_size
        
        metadata = {
            'format': file_path.suffix.lower(),
            'size_bytes': file_size,
            'processing_note': 'Audio transcription not implemented'
        }
        
        text = f"Audio file: {file_path.name}\n"
        text += f"Format: {file_path.suffix.upper()}\n"
        text += f"Size: {file_size} bytes"
        
        return {
            'text': text,
            'metadata': metadata,
            'content_type': 'audio'
        }
    
    async def _process_video(self, file_path: Path) -> Dict[str, Any]:
        """Process video files"""
        # Basic video file analysis
        file_size = file_path.stat().st_size
        
        metadata = {
            'format': file_path.suffix.lower(),
            'size_bytes': file_size,
            'processing_note': 'Video analysis not implemented'
        }
        
        text = f"Video file: {file_path.name}\n"
        text += f"Format: {file_path.suffix.upper()}\n"
        text += f"Size: {file_size} bytes"
        
        return {
            'text': text,
            'metadata': metadata,
            'content_type': 'video'
        }
    
    async def _generate_embeddings(self, file_record: FileRecord, persona_type: PersonaType):
        """Generate embeddings for file content"""
        try:
            text = file_record.extracted_text
            if not text or len(text.strip()) < 10:
                return
            
            # Split text into chunks
            chunks = self._split_text_into_chunks(text)
            
            if not chunks:
                return
            
            # Generate embeddings for each chunk
            embeddings = self.embedding_model.encode(chunks)
            
            # Store embeddings in vector database
            chunk_ids = []
            chunk_metadata = []
            
            for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = f"{file_record.id}_{i}"
                chunk_ids.append(chunk_id)
                
                metadata = {
                    'file_id': str(file_record.id),
                    'persona_type': persona_type.value,
                    'file_type': file_record.file_type,
                    'chunk_index': i,
                    'content': chunk_text[:500],  # First 500 chars for preview
                    'timestamp': datetime.utcnow().isoformat()
                }
                chunk_metadata.append(metadata)
            
            # Add to vector store
            success = await vector_store.add_vectors(
                collection_name="documents",
                vectors=embeddings.tolist(),
                metadata=chunk_metadata,
                ids=chunk_ids
            )
            
            if success:
                file_record.chunk_count = len(chunks)
                file_record.embedding_model = self.embedding_model_name
                file_record.embedding_dimensions = len(embeddings[0])
                logger.info(f"Generated embeddings for {file_record.original_filename}: {len(chunks)} chunks")
            
        except Exception as e:
            logger.error("Embedding generation failed", 
                        filename=file_record.original_filename, 
                        error=str(e))
            raise
    
    def _split_text_into_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for embedding"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence or paragraph boundary
            if end < len(text):
                # Look for sentence endings
                for boundary in ['. ', '.\n', '!\n', '?\n']:
                    boundary_pos = text.rfind(boundary, start, end)
                    if boundary_pos > start:
                        end = boundary_pos + len(boundary)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    async def _save_file(self, file_data: BinaryIO, filename: str) -> Path:
        """Save uploaded file to storage"""
        file_path = self.upload_dir / filename
        
        with open(file_path, 'wb') as f:
            while True:
                chunk = file_data.read(self.chunk_size)
                if not chunk:
                    break
                f.write(chunk)
        
        return file_path
    
    async def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate SHA-256 checksum of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _generate_unique_filename(self, original_filename: str) -> str:
        """Generate unique filename with timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name, ext = os.path.splitext(original_filename)
        return f"{timestamp}_{name}{ext}"
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type using python-magic"""
        try:
            mime_type = magic.from_file(str(file_path), mime=True)
            extension = file_path.suffix.lower()[1:]  # Remove the dot
            
            # Map MIME types to our file types
            type_mapping = {
                'application/pdf': 'pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                'text/plain': 'txt',
                'text/markdown': 'md',
                'text/html': 'html',
                'application/json': 'json',
                'text/csv': 'csv',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
                'application/zip': 'zip',
                'image/jpeg': 'jpg',
                'image/png': 'png',
                'audio/mpeg': 'mp3',
                'video/mp4': 'mp4'
            }
            
            # Try MIME type first, then extension
            return type_mapping.get(mime_type, extension or 'unknown')
            
        except Exception as e:
            logger.warning("File type detection failed", error=str(e))
            return file_path.suffix.lower()[1:] if file_path.suffix else 'unknown'

# Global file processor instance
file_processor = FileProcessor() 