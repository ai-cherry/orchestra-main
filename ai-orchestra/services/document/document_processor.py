"""
Document processor for large files in AI Orchestra.

This module provides functionality for processing large documents,
extracting text, and chunking them for efficient storage and retrieval.
"""

import json
import logging
import os
import re
import uuid
from typing import Any, Dict, List, Optional, Tuple

from ai_orchestra.core.config import settings

logger = logging.getLogger("ai_orchestra.services.document.document_processor")


class TextExtractionError(Exception):
    """Exception for text extraction-related errors."""


class ChunkingStrategy:
    """Chunking strategies for text documents."""

    PARAGRAPH = "paragraph"
    SENTENCE = "sentence"
    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"
    HYBRID = "hybrid"


class DocumentProcessor:
    """
    Document processor for large files.

    This class provides methods for extracting text from various file formats
    and chunking it into smaller pieces for efficient storage and retrieval.
    """

    def __init__(self, ai_service=None):
        """
        Initialize the document processor.

        Args:
            ai_service: Optional AI service for semantic operations
        """
        self.ai_service = ai_service
        self.chunk_size = settings.documents.chunk_size
        self.chunk_overlap = settings.documents.chunk_overlap
        self.max_file_size = settings.documents.max_file_size

    async def process_file(
        self,
        file_path: str,
        document_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        chunking_strategy: str = ChunkingStrategy.PARAGRAPH,
    ) -> Tuple[str, List[Dict[str, Any]], Dict[str, Any]]:
        """
        Process a file to extract text and create chunks.

        Args:
            file_path: Path to the file
            document_id: Optional document ID (generated if not provided)
            metadata: Optional metadata
            chunking_strategy: Chunking strategy to use

        Returns:
            Tuple of (document_id, chunks, document_metadata)

        Raises:
            TextExtractionError: If text extraction fails
            ValueError: If file is too large
        """
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise ValueError(
                f"File size ({file_size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
            )

        # Generate document ID if not provided
        if not document_id:
            document_id = str(uuid.uuid4())

        # Extract text from file
        text, file_type = await self._extract_text(file_path)

        # Get file metadata
        file_metadata = self._get_file_metadata(file_path, file_type)

        # Combine with provided metadata
        if metadata:
            file_metadata.update(metadata)

        # Create document metadata
        document_metadata = {
            "id": document_id,
            "title": os.path.basename(file_path),
            "source": file_path,
            "file_type": file_type,
            "file_size": file_size,
            "metadata": file_metadata,
        }

        # Chunk the text
        if chunking_strategy == ChunkingStrategy.PARAGRAPH:
            chunks = self._chunk_by_paragraph(text, document_id)
        elif chunking_strategy == ChunkingStrategy.SENTENCE:
            chunks = self._chunk_by_sentence(text, document_id)
        elif chunking_strategy == ChunkingStrategy.FIXED_SIZE:
            chunks = self._chunk_by_fixed_size(text, document_id)
        elif chunking_strategy == ChunkingStrategy.SEMANTIC:
            chunks = await self._chunk_by_semantic(text, document_id)
        elif chunking_strategy == ChunkingStrategy.HYBRID:
            chunks = await self._chunk_by_hybrid(text, document_id)
        else:
            # Default to paragraph chunking
            chunks = self._chunk_by_paragraph(text, document_id)

        return document_id, chunks, document_metadata

    async def _extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from a file based on its type.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (extracted text, detected file type)

        Raises:
            TextExtractionError: If extraction fails
        """
        # Detect file type based on extension
        _, extension = os.path.splitext(file_path.lower())
        extension = extension[1:] if extension else ""

        # Extract text based on file type
        try:
            if extension == "pdf":
                return await self._extract_from_pdf(file_path), "pdf"
            elif extension == "txt":
                return await self._extract_from_text(file_path), "txt"
            elif extension in ["md", "markdown"]:
                return await self._extract_from_text(file_path), "markdown"
            elif extension in ["docx", "doc"]:
                return await self._extract_from_docx(file_path), "docx"
            elif extension == "json":
                return await self._extract_from_json(file_path), "json"
            elif extension in ["csv", "tsv"]:
                return await self._extract_from_csv(file_path), "csv"
            elif extension in ["html", "htm"]:
                return await self._extract_from_html(file_path), "html"
            else:
                # Try as text file
                return await self._extract_from_text(file_path), "txt"
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise TextExtractionError(f"Failed to extract text: {e}")

    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            # Try PyMuPDF (fitz) first
            import fitz

            doc = fitz.open(file_path)
            text = ""

            # Process each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Add page number for context
                text += f"--- Page {page_num + 1} ---\n"
                text += page.get_text() + "\n\n"

            return text.strip()
        except ImportError:
            try:
                # Fall back to pypdf
                from pypdf import PdfReader

                reader = PdfReader(file_path)
                text = ""

                # Process each page
                for i, page in enumerate(reader.pages):
                    # Add page number for context
                    text += f"--- Page {i + 1} ---\n"
                    text += page.extract_text() + "\n\n"

                return text.strip()
            except ImportError:
                raise TextExtractionError(
                    "PDF extraction requires either PyMuPDF or pypdf. "
                    "Install with: pip install pymupdf or pip install pypdf"
                )

    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from a text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            # If all attempts fail, use binary mode and decode with errors ignored
            with open(file_path, "rb") as f:
                return f.read().decode("utf-8", errors="ignore")

    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            import docx

            doc = docx.Document(file_path)

            # Extract text from paragraphs
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"

            return text.strip()
        except ImportError:
            raise TextExtractionError(
                "DOCX extraction requires python-docx. "
                "Install with: pip install python-docx"
            )

    async def _extract_from_json(self, file_path: str) -> str:
        """Extract text from a JSON file."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

            # Convert JSON to formatted string
            return json.dumps(data, indent=2)

    async def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from a CSV file."""
        import csv

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)

            # Get header
            header = next(reader, [])

            # Convert rows to formatted text
            text = "| " + " | ".join(header) + " |\n"
            text += "| " + " | ".join(["---"] * len(header)) + " |\n"

            for row in reader:
                text += "| " + " | ".join(row) + " |\n"

            return text.strip()

    async def _extract_from_html(self, file_path: str) -> str:
        """Extract text from an HTML file."""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")

                # Extract text
                return soup.get_text(separator="\n", strip=True)
        except ImportError:
            raise TextExtractionError(
                "HTML extraction requires beautifulsoup4. "
                "Install with: pip install beautifulsoup4"
            )

    def _get_file_metadata(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Get metadata for a file.

        Args:
            file_path: Path to the file
            file_type: Type of the file

        Returns:
            File metadata
        """
        file_stats = os.stat(file_path)

        return {
            "filename": os.path.basename(file_path),
            "file_path": file_path,
            "file_type": file_type,
            "file_size": file_stats.st_size,
            "created_at": file_stats.st_ctime,
            "modified_at": file_stats.st_mtime,
        }

    def _chunk_by_paragraph(self, text: str, document_id: str) -> List[Dict[str, Any]]:
        """
        Split content by paragraphs.

        Args:
            text: Text content to chunk
            document_id: Document ID

        Returns:
            List of chunks with metadata
        """
        # Split by double newlines (paragraphs)
        paragraphs = re.split(r"\n\s*\n", text)

        # Filter out empty paragraphs
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        # Create chunks
        chunks = []
        current_chunk = []
        current_length = 0

        for i, paragraph in enumerate(paragraphs):
            # If adding this paragraph would exceed max length and we have content,
            # create a new chunk
            if current_length + len(paragraph) > self.chunk_size and current_chunk:
                # Join current chunks and add to result
                chunk_text = "\n\n".join(current_chunk)
                chunk_id = f"{document_id}_chunk_{len(chunks)}"

                chunks.append(
                    {
                        "id": chunk_id,
                        "document_id": document_id,
                        "content": chunk_text,
                        "metadata": {
                            "seq_num": len(chunks),
                            "chunk_type": "paragraph",
                            "start_idx": i - len(current_chunk),
                            "end_idx": i - 1,
                        },
                    }
                )

                # Start new chunk, possibly with overlap
                if self.chunk_overlap > 0 and len(current_chunk) > 1:
                    # Keep the last paragraph for overlap
                    current_chunk = current_chunk[-1:]
                    current_length = len(current_chunk[0])
                else:
                    current_chunk = []
                    current_length = 0

            # Add paragraph to current chunk
            current_chunk.append(paragraph)
            current_length += len(paragraph)

        # Add the last chunk if not empty
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            chunk_id = f"{document_id}_chunk_{len(chunks)}"

            chunks.append(
                {
                    "id": chunk_id,
                    "document_id": document_id,
                    "content": chunk_text,
                    "metadata": {
                        "seq_num": len(chunks),
                        "chunk_type": "paragraph",
                        "start_idx": len(paragraphs) - len(current_chunk),
                        "end_idx": len(paragraphs) - 1,
                    },
                }
            )

        return chunks

    def _chunk_by_sentence(self, text: str, document_id: str) -> List[Dict[str, Any]]:
        """
        Split content by sentences.

        Args:
            text: Text content to chunk
            document_id: Document ID

        Returns:
            List of chunks with metadata
        """
        # Split by sentences
        sentences = re.split(r"(?<=[.!?])\s+", text)

        # Filter out empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]

        # Create chunks
        chunks = []
        current_chunk = []
        current_length = 0

        for i, sentence in enumerate(sentences):
            # If adding this sentence would exceed max length and we have content,
            # create a new chunk
            if current_length + len(sentence) > self.chunk_size and current_chunk:
                # Join current chunks and add to result
                chunk_text = " ".join(current_chunk)
                chunk_id = f"{document_id}_chunk_{len(chunks)}"

                chunks.append(
                    {
                        "id": chunk_id,
                        "document_id": document_id,
                        "content": chunk_text,
                        "metadata": {
                            "seq_num": len(chunks),
                            "chunk_type": "sentence",
                            "start_idx": i - len(current_chunk),
                            "end_idx": i - 1,
                        },
                    }
                )

                # Start new chunk, possibly with overlap
                if self.chunk_overlap > 0 and len(current_chunk) > 1:
                    # Keep the last few sentences for overlap
                    overlap_sentences = min(
                        int(self.chunk_overlap / 100 * len(current_chunk)),
                        len(current_chunk),
                    )
                    current_chunk = current_chunk[-overlap_sentences:]
                    current_length = sum(len(s) for s in current_chunk)
                else:
                    current_chunk = []
                    current_length = 0

            # Add sentence to current chunk
            current_chunk.append(sentence)
            current_length += len(sentence)

        # Add the last chunk if not empty
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk_id = f"{document_id}_chunk_{len(chunks)}"

            chunks.append(
                {
                    "id": chunk_id,
                    "document_id": document_id,
                    "content": chunk_text,
                    "metadata": {
                        "seq_num": len(chunks),
                        "chunk_type": "sentence",
                        "start_idx": len(sentences) - len(current_chunk),
                        "end_idx": len(sentences) - 1,
                    },
                }
            )

        return chunks

    def _chunk_by_fixed_size(self, text: str, document_id: str) -> List[Dict[str, Any]]:
        """
        Split content by fixed size chunks.

        Args:
            text: Text content to chunk
            document_id: Document ID

        Returns:
            List of chunks with metadata
        """
        chunks = []

        # Create chunks with overlap
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            # Get chunk content
            chunk_end = min(i + self.chunk_size, len(text))
            chunk_text = text[i:chunk_end]

            # Skip if chunk is too small
            if len(chunk_text) < 100:  # Minimum meaningful chunk size
                continue

            # Create chunk
            chunk_id = f"{document_id}_chunk_{len(chunks)}"
            chunks.append(
                {
                    "id": chunk_id,
                    "document_id": document_id,
                    "content": chunk_text,
                    "metadata": {
                        "seq_num": len(chunks),
                        "chunk_type": "fixed_size",
                        "start_char": i,
                        "end_char": chunk_end,
                    },
                }
            )

        return chunks

    async def _chunk_by_semantic(
        self, text: str, document_id: str
    ) -> List[Dict[str, Any]]:
        """
        Split content by semantic boundaries using AI.

        Args:
            text: Text content to chunk
            document_id: Document ID

        Returns:
            List of chunks with metadata
        """
        # Check if AI service is available
        if not self.ai_service:
            logger.warning(
                "AI service not available for semantic chunking, falling back to paragraph chunking"
            )
            return self._chunk_by_paragraph(text, document_id)

        try:
            # Create prompt for semantic chunking
            prompt = f"""
            Split the following text into coherent semantic chunks. Each chunk should be a self-contained unit of information.

            Rules for chunking:
            1. Each chunk should be coherent and self-contained
            2. Preserve paragraph boundaries when possible
            3. Keep related information together
            4. Split at natural semantic boundaries
            5. Aim for chunks of roughly {self.chunk_size} characters

            Format your response as a JSON array of chunks, where each chunk is a string.

            Text to chunk:
            {text[:20000]}  # Limit text size for prompt

            Chunks:
            """

            # Generate chunks using AI
            semantic_chunks_text = await self.ai_service.generate_text(
                prompt=prompt, max_tokens=4096, temperature=0.2
            )

            # Parse JSON response
            try:
                # Extract JSON array from response
                json_match = re.search(r"\[.*\]", semantic_chunks_text, re.DOTALL)
                if json_match:
                    semantic_chunks = json.loads(json_match.group(0))
                else:
                    raise ValueError("No JSON array found in response")

                # Create chunks
                chunks = []
                for i, chunk_text in enumerate(semantic_chunks):
                    chunk_id = f"{document_id}_chunk_{i}"
                    chunks.append(
                        {
                            "id": chunk_id,
                            "document_id": document_id,
                            "content": chunk_text.strip(),
                            "metadata": {"seq_num": i, "chunk_type": "semantic"},
                        }
                    )

                return chunks
            except (json.JSONDecodeError, ValueError) as e:
                logger.error(f"Failed to parse semantic chunks: {e}")
                return self._chunk_by_paragraph(text, document_id)

        except Exception as e:
            logger.error(f"Semantic chunking failed: {e}")
            return self._chunk_by_paragraph(text, document_id)

    async def _chunk_by_hybrid(
        self, text: str, document_id: str
    ) -> List[Dict[str, Any]]:
        """
        Split content using a hybrid of multiple strategies.

        Args:
            text: Text content to chunk
            document_id: Document ID

        Returns:
            List of chunks with metadata
        """
        # First split by paragraphs
        paragraph_chunks = self._chunk_by_paragraph(text, document_id)

        # For large paragraphs, split further by sentences
        refined_chunks = []
        for i, chunk in enumerate(paragraph_chunks):
            if len(chunk["content"]) > self.chunk_size / 2:
                # Create a temporary document ID for sub-chunking
                temp_id = f"{document_id}_sub_{i}"

                # Split by sentences
                sub_chunks = self._chunk_by_sentence(chunk["content"], temp_id)

                # Add refined chunks with updated metadata
                for j, sub_chunk in enumerate(sub_chunks):
                    sub_chunk_id = f"{document_id}_chunk_{len(refined_chunks)}"
                    refined_chunks.append(
                        {
                            "id": sub_chunk_id,
                            "document_id": document_id,
                            "content": sub_chunk["content"],
                            "metadata": {
                                "seq_num": len(refined_chunks),
                                "chunk_type": "hybrid",
                                "parent_chunk": i,
                                "sub_chunk": j,
                            },
                        }
                    )
            else:
                # Keep chunk as is, just update the ID
                chunk["id"] = f"{document_id}_chunk_{len(refined_chunks)}"
                chunk["metadata"]["seq_num"] = len(refined_chunks)
                chunk["metadata"]["chunk_type"] = "hybrid"
                refined_chunks.append(chunk)

        return refined_chunks
