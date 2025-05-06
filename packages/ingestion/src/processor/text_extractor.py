"""
Text Extraction Module for File Ingestion System.

This module provides functionality for extracting text content from 
various file formats such as PDF, DOCX, XLSX, etc.
"""

import os
import logging
import tempfile
from typing import Dict, List, Optional, Any, Tuple, BinaryIO, Union
import mimetypes
import json
import csv
import io
import zipfile
import tarfile
import re

from packages.ingestion.src.models.ingestion_models import FileType
from packages.ingestion.src.config.settings import get_settings

# Configure logging
logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Exception for text extraction-related errors."""
    pass


class TextExtractor:
    """
    Text extraction implementation for various file types.
    
    This class provides methods for extracting text content from
    different file formats using appropriate libraries.
    """
    
    def __init__(self):
        """Initialize the text extractor with settings."""
        self.settings = get_settings()
        
    async def extract_text(self, file_path: str) -> Tuple[str, FileType]:
        """
        Extract text from a file based on its type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted text, detected file type)
            
        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            # Detect file type
            file_type = self._detect_file_type(file_path)
            
            # Extract text based on file type
            if file_type == FileType.PDF:
                text = await self._extract_from_pdf(file_path)
            elif file_type == FileType.DOCX:
                text = await self._extract_from_docx(file_path)
            elif file_type == FileType.XLSX:
                text = await self._extract_from_xlsx(file_path)
            elif file_type == FileType.CSV:
                text = await self._extract_from_csv(file_path)
            elif file_type in [FileType.TXT, FileType.JSON, FileType.HTML]:
                text = await self._extract_from_text_file(file_path)
            elif file_type == FileType.ZIP:
                # For archives, we should unpack and process each file separately
                # Here we'll just extract a summary for demonstration
                text = await self._summarize_archive(file_path, is_zip=True)
            elif file_type == FileType.TAR_GZ:
                # Same for tar.gz archives
                text = await self._summarize_archive(file_path, is_zip=False)
            elif file_type in [FileType.JPG, FileType.PNG]:
                text = await self._extract_from_image(file_path)
            elif file_type == FileType.EML:
                text = await self._extract_from_email(file_path)
            elif file_type == FileType.MSG:
                text = await self._extract_from_msg(file_path)
            else:
                # Generic text extraction for unsupported types
                text = await self._extract_generic(file_path)
                
            return text, file_type
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise TextExtractionError(f"Failed to extract text: {e}")
    
    def _detect_file_type(self, file_path: str) -> FileType:
        """
        Detect file type based on extension and content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Detected file type
        """
        # Get file extension
        _, extension = os.path.splitext(file_path.lower())
        extension = extension[1:] if extension else ""
        
        # Check extension against known types
        if extension == "pdf":
            return FileType.PDF
        elif extension == "docx":
            return FileType.DOCX
        elif extension == "xlsx":
            return FileType.XLSX
        elif extension == "csv":
            return FileType.CSV
        elif extension == "txt":
            return FileType.TXT
        elif extension == "json":
            return FileType.JSON
        elif extension == "html" or extension == "htm":
            return FileType.HTML
        elif extension == "zip":
            return FileType.ZIP
        elif extension in ["tar.gz", "tgz"]:
            return FileType.TAR_GZ
        elif extension in ["jpg", "jpeg"]:
            return FileType.JPG
        elif extension == "png":
            return FileType.PNG
        elif extension == "eml":
            return FileType.EML
        elif extension == "msg":
            return FileType.MSG
        else:
            # Try to guess type from content
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                if mime_type == "application/pdf":
                    return FileType.PDF
                elif "officedocument.wordprocessingml" in mime_type:
                    return FileType.DOCX
                elif "officedocument.spreadsheetml" in mime_type:
                    return FileType.XLSX
                elif mime_type == "text/csv":
                    return FileType.CSV
                elif mime_type == "text/plain":
                    return FileType.TXT
                elif mime_type == "application/json":
                    return FileType.JSON
                elif mime_type.startswith("text/html"):
                    return FileType.HTML
                elif mime_type == "application/zip":
                    return FileType.ZIP
                elif mime_type in ["application/x-tar", "application/gzip"]:
                    return FileType.TAR_GZ
                elif mime_type.startswith("image/jpeg"):
                    return FileType.JPG
                elif mime_type == "image/png":
                    return FileType.PNG
                elif mime_type == "message/rfc822":
                    return FileType.EML
                elif mime_type == "application/vnd.ms-outlook":
                    return FileType.MSG
            
            logger.warning(f"Unknown file type for {file_path}, treating as plain text")
            return FileType.UNKNOWN
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            # Import PyMuPDF (fitz) for PDF extraction
            import fitz
        except ImportError:
            try:
                # Fall back to pypdf if fitz not available
                from pypdf import PdfReader
                
                with open(file_path, "rb") as f:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n\n"
                    
                return text.strip()
            except ImportError:
                raise TextExtractionError(
                    "PDF extraction requires either PyMuPDF or pypdf. "
                    "Install with: pip install pypdf or pip install pymupdf"
                )
        
        # Extract using PyMuPDF (more robust)
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text() + "\n\n"
            
        return text.strip()
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            # Import docx library
            import docx
        except ImportError:
            raise TextExtractionError(
                "DOCX extraction requires python-docx. "
                "Install with: pip install python-docx"
            )
        
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        
        return text.strip()
    
    async def _extract_from_xlsx(self, file_path: str) -> str:
        """Extract text from an XLSX file."""
        try:
            # Import pandas for Excel processing
            import pandas as pd
        except ImportError:
            try:
                # Fall back to openpyxl if pandas not available
                import openpyxl
                
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                text = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"Sheet: {sheet_name}\n\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        # Convert all values to string and join
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        text += row_text + "\n"
                    
                    text += "\n\n"
                
                return text.strip()
            except ImportError:
                raise TextExtractionError(
                    "XLSX extraction requires either pandas or openpyxl. "
                    "Install with: pip install pandas openpyxl"
                )
        
        # Use pandas to read Excel file (more robust)
        dfs = pd.read_excel(file_path, sheet_name=None)
        text = ""
        
        for sheet_name, df in dfs.items():
            text += f"Sheet: {sheet_name}\n\n"
            text += df.to_string() + "\n\n"
            
        return text.strip()
    
    async def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from a CSV file."""
        text = ""
        
        with open(file_path, "r", newline="", encoding="utf-8", errors="ignore") as csvfile:
            try:
                # Try to detect dialect
                dialect = csv.Sniffer().sniff(csvfile.read(1024))
                csvfile.seek(0)
                
                reader = csv.reader(csvfile, dialect)
                for row in reader:
                    text += " | ".join(row) + "\n"
            except:
                # If dialect detection fails, try basic CSV parsing
                csvfile.seek(0)
                reader = csv.reader(csvfile)
                for row in reader:
                    text += " | ".join(row) + "\n"
        
        return text.strip()
    
    async def _extract_from_text_file(self, file_path: str) -> str:
        """Extract text from a plain text, JSON, or HTML file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
        return text
    
    async def _summarize_archive(self, file_path: str, is_zip: bool = True) -> str:
        """
        Summarize contents of an archive file.
        
        For full processing, archives should be extracted and each file
        processed separately. This method just provides a summary.
        """
        if is_zip:
            # Handle ZIP archives
            with zipfile.ZipFile(file_path, "r") as archive:
                file_list = archive.namelist()
        else:
            # Handle TAR.GZ archives
            with tarfile.open(file_path, "r:*") as archive:
                file_list = archive.getnames()
                
        # Create a summary
        file_count = len(file_list)
        summary_text = f"Archive with {file_count} files:\n\n"
        for i, filename in enumerate(file_list[:100]):  # Limit to first 100 files
            summary_text += f"- {filename}\n"
            
        if file_count > 100:
            summary_text += f"\n... and {file_count - 100} more files."
            
        return summary_text
    
    async def _extract_from_image(self, file_path: str) -> str:
        """
        Extract text from an image file (OCR).
        
        This requires Google Cloud Vision API or another OCR solution.
        """
        try:
            from google.cloud import vision
        except ImportError:
            return f"[Image file: {os.path.basename(file_path)}. OCR not performed - requires google-cloud-vision package.]"
        
        try:
            # Create a client
            client = vision.ImageAnnotatorClient()
            
            # Load image
            with open(file_path, "rb") as image_file:
                content = image_file.read()
                
            image = vision.Image(content=content)
            
            # Perform text detection
            response = client.text_detection(image=image)
            texts = response.text_annotations
            
            if texts:
                # First result is the complete text
                return texts[0].description
            else:
                return f"[Image file: {os.path.basename(file_path)}. No text detected.]"
        except Exception as e:
            logger.error(f"Error during OCR of image {file_path}: {e}")
            return f"[Image file: {os.path.basename(file_path)}. OCR failed: {str(e)}]"
    
    async def _extract_from_email(self, file_path: str) -> str:
        """Extract text from an EML email file."""
        try:
            import email
            from email.policy import default
        except ImportError:
            raise TextExtractionError("Email extraction requires the email package.")
            
        with open(file_path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=default)
            
        # Extract headers
        text = f"From: {msg.get('From', '')}\n"
        text += f"To: {msg.get('To', '')}\n"
        text += f"Subject: {msg.get('Subject', '')}\n"
        text += f"Date: {msg.get('Date', '')}\n\n"
        
        # Extract content
        if msg.is_multipart():
            for part in msg.iter_parts():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    text += part.get_content()
                elif content_type == "text/html":
                    # For HTML parts, we could extract text or include a note
                    text += f"\n[HTML Content: {len(part.get_content())} bytes]\n"
        else:
            if msg.get_content_type() == "text/plain":
                text += msg.get_content()
            elif msg.get_content_type() == "text/html":
                # Basic HTML tag removal - a more robust approach would use html2text or BeautifulSoup
                html_content = msg.get_content()
                text += re.sub(r'<[^>]+>', '', html_content)
                
        return text
    
    async def _extract_from_msg(self, file_path: str) -> str:
        """Extract text from an Outlook MSG file."""
        try:
            import extract_msg
        except ImportError:
            return (
                f"[MSG file: {os.path.basename(file_path)}. "
                f"Extraction not performed - requires extract_msg package.]"
            )
        
        try:
            msg = extract_msg.Message(file_path)
            
            # Extract headers
            text = f"From: {msg.sender}\n"
            text += f"To: {msg.to}\n"
            text += f"Subject: {msg.subject}\n"
            text += f"Date: {msg.date}\n\n"
            
            # Extract body
            text += msg.body
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from MSG file {file_path}: {e}")
            return f"[MSG file: {os.path.basename(file_path)}. Extraction failed: {str(e)}]"
    
    async def _extract_generic(self, file_path: str) -> str:
        """
        Generic text extraction for unsupported file types.
        
        This method attempts to read the file as text and extracts
        any visible text strings.
        """
        try:
            # Try to read as text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
                
            # If text seems meaningful (not binary), return it
            if self._is_meaningful_text(text):
                return text
                
            # Otherwise, try binary mode and extract strings
            return f"[Binary file: {os.path.basename(file_path)}. No text extraction available.]"
        except Exception as e:
            logger.error(f"Error in generic text extraction for {file_path}: {e}")
            return f"[File: {os.path.basename(file_path)}. Text extraction not supported.]"
    
    def _is_meaningful_text(self, text: str, threshold: float = 0.7) -> bool:
        """
        Check if text content appears to be meaningful text rather than binary.
        
        Args:
            text: Text content to check
            threshold: Threshold of printable characters (0.0 to 1.0)
            
        Returns:
            True if text appears to be meaningful
        """
        if not text:
            return False
            
        # Count printable ASCII characters
        printable_count = sum(c.isprintable() for c in text)
        
        # Calculate ratio of printable characters
        printable_ratio = printable_count / len(text)
        
        return printable_ratio > threshold
        
    async def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to split
            chunk_size: Maximum size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
            
        # If text is shorter than chunk_size, return as is
        if len(text) <= chunk_size:
            return [text]
            
        chunks = []
        start = 0
        
        while start < len(text):
            # Get chunk with specified size
            end = start + chunk_size
            
            # Adjust end to avoid splitting words
            if end < len(text):
                # Try to find a good split point (space, newline)
                good_end = end
                while good_end > start + chunk_size - 50:
                    if text[good_end] in [" ", "\n", ".", ",", ";", ":", "!", "?"]:
                        break
                    good_end -= 1
                    
                if good_end > start + chunk_size - 50:
                    end = good_end + 1
            
            # Add chunk
            chunks.append(text[start:end])
            
            # Move start position for next chunk
            start = end - overlap
            
        return chunks
